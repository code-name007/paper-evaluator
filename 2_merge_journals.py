#!/usr/bin/env python3
"""
合并DOCX + XLSX，生成高水平期刊完整目录
"""
from docx import Document
import openpyxl
import json
import re

# ====== 第一部分：从DOCX读取1-4类国际期刊 ======
doc = Document('/Users/code-007/Downloads/高水平五类期刊目录/高水平期刊目录-正式发文-20230919.docx')

TIER_NAMES = {
    0: ('高水平1类', 0),   # 国际顶级
    1: ('高水平2类', 1),   # 国际一流
    2: ('高水平3类', 2),   # 国际高水平
    3: ('高水平4类', 3),   # 国际重要
}

journals = {}

for ti, tbl in enumerate(doc.tables):
    if ti > 3:  # 表格5和6是中文期刊，单独处理
        continue
    tier_name, tier_level = TIER_NAMES[ti]
    for ri, row in enumerate(tbl.rows):
        if ri == 0: continue
        cells = [c.text.strip() for c in row.cells]
        if not cells[0]: continue
        name = cells[0]
        issn = cells[1] if len(cells) > 1 else ''
        eissn = cells[2] if len(cells) > 2 else ''
        subject = cells[3] if len(cells) > 3 else ''
        for key in [name.lower(), issn, eissn]:
            if key:
                journals[key] = {
                    'name': name, 'issn': issn, 'eissn': eissn,
                    'tier': tier_name, 'tier_level': tier_level,
                    'subject': subject, 'source': 'docx'
                }
        if issn:
            journals[issn] = journals[name.lower()]
        if eissn:
            journals[eissn] = journals[name.lower()]

# 表格5和6：中文期刊（5类）
for ti, tbl in enumerate(doc.tables[4:], start=4):
    for ri, row in enumerate(tbl.rows):
        if ri == 0: continue
        cells = [c.text.strip() for c in row.cells]
        if not cells[0]: continue
        name = cells[0]
        issn = cells[1] if len(cells) > 1 else ''
        tier_name = '高水平5A类' if ti == 4 else '高水平5B类'
        tier_level = 4 if ti == 4 else 5
        key = issn if issn else name.lower()
        journals[key] = {
            'name': name, 'issn': issn, 'eissn': '',
            'tier': tier_name, 'tier_level': tier_level,
            'subject': '', 'source': 'docx'
        }

print(f"DOCX期刊: {len(set(k for k in journals if len(k)>3))}种 (1-5类)")

# ====== 第二部分：从XLSX读取各学院5A/5B期刊 ======
wb = openpyxl.load_workbook('/Users/code-007/Downloads/高水平五类期刊目录/国际高质量期刊五类目录汇总表-终稿-20231008.xlsx')
ws = wb['总表']

current_tier = None
xlsx_added = 0

for ri in range(1, ws.max_row + 1):
    col2 = str(ws.cell(ri, 2).value or '').strip()
    col3 = str(ws.cell(ri, 3).value or '').strip()

    if 'A档' in col2:
        current_tier = ('高水平5A类', 4)
        continue
    elif 'B档' in col2:
        current_tier = ('高水平5B类', 5)
        continue
    elif not col2.isdigit() or not col3:
        continue

    name = col3
    issn = str(ws.cell(ri, 4).value or '').strip()
    eissn = str(ws.cell(ri, 5).value or '').strip()
    subject = str(ws.cell(ri, 6).value or '').strip()
    zone = str(ws.cell(ri, 7).value or '').strip()

    if not current_tier: continue

    tier_name, tier_level = current_tier
    # 用ISSN去重，已有的不覆盖（DOCX优先，因为它更权威）
    for lookup in [issn, eissn, name.lower()]:
        if lookup and lookup in journals:
            break
    else:
        key = issn if issn else name.lower()
        journals[key] = {
            'name': name, 'issn': issn, 'eissn': eissn,
            'tier': tier_name, 'tier_level': tier_level,
            'subject': subject, 'zone': zone,
            'source': 'xlsx'
        }
        xlsx_added += 1

print(f"XLSX新增期刊: {xlsx_added}种 (5A/5B类)")
print(f"合并后总计: {len(journals)}条记录")

# ====== 保存 ======
out = '/Users/code-007/Downloads/啵啵龙/文稿/论文评价网站/journals_data.json'
with open(out, 'w', encoding='utf-8') as f:
    json.dump(journals, f, ensure_ascii=False, indent=2)

# 统计
from collections import Counter
tier_stats = Counter(v['tier'] for v in journals.values())
for tier in sorted(tier_stats, key=lambda x: journals.get(x, {}).get('tier_level', 9)):
    pass
print("\n各层级期刊数量:")
for k, v in sorted(tier_stats.items(), key=lambda x: x[0]):
    print(f"  {k}: {v}")

print(f"\n已保存: {out}")
