#!/usr/bin/env python3
"""
论文情况评价 & 高水平期刊匹配系统
哈尔滨医科大学人事处专用
"""
import streamlit as st
import json
import re
import time
import pandas as pd
import requests
from pathlib import Path
import io
import sys

# ===== PDF 扫描件识别（增强版） =====
def preprocess_for_ocr(img):
    """
    对图像进行预处理以提升 OCR 准确率：
    1. 转为灰度
    2. 增加对比度
    3. 锐化
    4. 可选：降噪
    """
    from PIL import ImageEnhance, ImageFilter
    # 转灰度
    gray = img.convert('L')
    # 增加对比度（1.5倍）
    enh = ImageEnhance.Contrast(gray)
    enhanced = enh.enhance(1.5)
    # 锐化
    sharpened = enhanced.filter(ImageFilter.SHARPEN)
    # 边缘增强
    sharpened2 = sharpened.filter(ImageFilter.EDGE_ENHANCE)
    return sharpened2


# ===== PDF 文字提取（三策略并行，自动选最优） =====
def extract_text_from_pdf_v2(file_bytes, filename=""):
    """
    从 PDF 中提取文字，三策略并行，取最优结果：
    策略1：pdfplumber（原生，适合简单排版）
    策略2：PyMuPDF 直接提取（适合复杂排版、多栏）
    策略3：OCR（扫描件 / 前两者失败时）

    返回：(最佳文本, 是否为扫描件, 各策略文本长度)
    """
    import tempfile, os, io

    strategies = {}  # name -> (text, char_count)

    # ---- 策略1：pdfplumber ----
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            parts = []
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
        txt = "\n".join(parts)
        strategies["pdfplumber"] = (txt, len(txt.strip()))
    except Exception as e:
        strategies["pdfplumber"] = ("", 0)

    # ---- 策略2：PyMuPDF 直接提取（更好的排版处理）----
    tmp_path = None
    try:
        import fitz
        tmp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        tmp.write(file_bytes)
        tmp.close()
        tmp_path = tmp.name

        with fitz.open(tmp_path) as doc:
            parts = []
            for page in doc:
                # PyMuPDF 的 extractUSD 会保留阅读顺序，比 extract_text 更准
                blocks = page.get_text("blocks")
                # 按 (x0, y0) 排序：从上到下、从左到右
                blocks.sort(key=lambda b: (round(b[1] / 20) * 20, b[0]))
                page_text = ""
                for block in blocks:
                    block_text = block[4].strip()
                    if block_text:
                        page_text += block_text + "\n"
                if page_text.strip():
                    parts.append(page_text)
        txt = "\n".join(parts)
        strategies["pymupdf"] = (txt, len(txt.strip()))
    except Exception as e:
        strategies["pymupdf"] = ("", 0)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

    # ---- 选择最佳策略 ----
    # 优先：字符数 > 50 且最连贯的策略
    best_name = None
    best_len = 0
    best_text = ""

    for name, (txt, clen) in strategies.items():
        if clen > best_len:
            best_len = clen
            best_text = txt
            best_name = name

    # ---- 策略3：OCR（当前两者都失败或字符数太少时）----
    is_scanned = False
    if best_len < 200:
        is_scanned = True
        tmp_path = None
        try:
            import fitz, pytesseract
            from PIL import Image, ImageEnhance, ImageFilter

            tmp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
            tmp.write(file_bytes)
            tmp.close()
            tmp_path = tmp.name

            with fitz.open(tmp_path) as doc:
                ocr_pages = []
                for page in doc:
                    # 300 DPI
                    mat = fitz.Matrix(300/72, 300/72)
                    pix = page.get_pixmap(matrix=mat)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    # 预处理
                    gray = img.convert('L')
                    enh = ImageEnhance.Contrast(gray)
                    img_p = enh.enhance(1.5).filter(ImageFilter.SHARPEN)
                    # OCR
                    ocr_text = pytesseract.image_to_string(
                        img_p, lang='chi_sim+eng', config='--psm 3 --oem 3'
                    )
                    if ocr_text.strip():
                        ocr_pages.append(ocr_text.strip())
                best_text = "\n".join(ocr_pages)
                best_name = "ocr"
        except Exception as e:
            sys.stderr.write(f"OCR failed: {e}\n")
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    return best_text, is_scanned, {k: v[1] for k, v in strategies.items()}



st.set_page_config(
    page_title="论文评价系统 | 哈医大人事处",
    page_icon="📖",
    layout="wide"
)

# ===== 加载期刊数据 =====
@st.cache_data


# ===== PDF 转 Word（适合排版复杂的文本型 PDF） =====
def convert_pdf_to_docx_text(pdf_bytes):
    """
    将 PDF 转换为 docx，再读取纯文本。
    适合：多栏排版、特殊字体、复杂图表的 PDF。
    返回：(文本内容, 转换是否成功)
    """
    from pdf2docx import Converter
    from docx import Document
    import tempfile, os

    tmp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
    tmp_pdf.write(pdf_bytes)
    tmp_pdf.close()

    tmp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp_docx.close()
    docx_path = tmp_docx.name

    try:
        cv = Converter(tmp_pdf.name)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        doc = Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n'.join(paragraphs), True
    except Exception as e:
        sys.stderr.write(f"PDF转Word失败: {e}\n")
        return "", False
    finally:
        for p in [tmp_pdf.name, docx_path]:
            if os.path.exists(p):
                try: os.unlink(p)
                except: pass


# ===== EasyOCR 深度识别（扫描件增强 OCR） =====
def extract_text_by_easyocr(pdf_bytes, dpi=200):
    """
    使用 EasyOCR（深度学习 OCR引擎）对 PDF 进行识别。
    对中文识别率显著高于 Tesseract。
    首次运行需下载模型（约 2GB），已安装到本地后离线可用。
    返回：(识别的文本, 是否成功)
    """
    import easyocr
    import fitz
    from PIL import Image, ImageEnhance, ImageFilter
    import numpy as np
    import tempfile, os

    tmp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
    tmp_pdf.write(pdf_bytes)
    tmp_pdf.close()
    tmp_path = tmp_pdf.name

    try:
        reader = easyocr.Reader(['ch_sim', 'en'], gpu=False, verbose=False)
        with fitz.open(tmp_path) as doc:
            all_text = []
            for page in doc:
                mat = fitz.Matrix(dpi/72, dpi/72)
                pix = page.get_pixmap(matrix=mat)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                gray = img.convert('L')
                enh = ImageEnhance.Contrast(gray)
                img_p = enh.enhance(1.5).filter(ImageFilter.SHARPEN)
                results = reader.readtext(np.array(img_p), detail=0)
                page_text = ' '.join([t.strip() for t in results if t.strip() and len(t.strip()) > 1])
                if page_text:
                    all_text.append(page_text)
        return '\n'.join(all_text), True
    except Exception as e:
        sys.stderr.write(f"EasyOCR 识别失败: {e}\n")
        return "", False
    finally:
        if os.path.exists(tmp_path):
            try: os.unlink(tmp_path)
            except: pass


def load_journals():
    path = Path(__file__).parent / "journals_data.json"
    with open(path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    # 按 name 建索引
    name_map = {}
    issn_map = {}
    for k, v in data.items():
        if len(k) > 5:
            name_map[k.lower().strip()] = v
        elif '-' in k or k.replace(' ', '').isdecimal():
            issn_map[k.strip()] = v
    return name_map, issn_map

TIER_COLORS = {
    '高水平1类': '🔴',
    '高水平2类': '🟠',
    '高水平3类': '🟡',
    '高水平4类': '🟢',
    '高水平5A类': '🔵',
    '高水平5B类': '🔵',
}

TIER_BG = {
    '高水平1类': '#ffe5e5',
    '高水平2类': '#fff3e0',
    '高水平3类': '#fffde7',
    '高水平4类': '#e8f5e9',
    '高水平5A类': '#e3f2fd',
    '高水平5B类': '#e3f2fd',
}

def get_tier_label(v):
    return v['tier'] if v else None

def match_journal(journal_name, name_map, issn_map, session):
    """匹配期刊，返回匹配结果"""
    name_clean = journal_name.lower().strip().replace(' ', '')
    # 精确匹配
    if name_clean in name_map:
        return name_map[name_clean]
    # 模糊匹配（包含）
    for key, val in name_map.items():
        if name_clean in key or key in name_clean:
            return val
    # 去掉冠词再试
    for alt in [name_clean.replace('the',''), name_clean.replace('-',' '), name_clean.replace(':','')]:
        if alt in name_map:
            return name_map[alt]
    return None

def get_sci_zone_from_crossref(journal_name, session):
    """通过CrossRef API获取期刊Impact Factor，推断SCI分区"""
    try:
        url = 'https://api.crossref.org/journals'
        params = {'query': journal_name, 'rows': 3}
        resp = session.get(url, params=params, timeout=8)
        if resp.status_code == 200:
            items = resp.json().get('message', {}).get('items', [])
            for item in items:
                if item.get('title'):
                    return {
                        'journal': item['title'][0] if item.get('title') else journal_name,
                        'doi': item.get('DOI',''),
                        'if': item.get('metrics', {}).get(' cites-per-doc', 0),
                    }
    except Exception:
        pass
    return None

def normalize_ocr_text(text):
    """标准化 OCR 文本，去除 CJK 字符间随机空格，修复常见 OCR 错误"""
    if not text:
        return text

    # 1. 处理 [J] [1] [A] 等标记中的多余空格
    text = re.sub(r'\[\s*J\s*\]', '[J]', text)
    text = re.sub(r'\[\s*(\d+)\s*\]', r'[\1]', text)
    text = re.sub(r'\[\s*([A-Za-z])\s*\]', r'[\1]', text)

    # 2. 去除 CJK 字符之间的空格
    # 用占位符保护 [J]
    marker_J = '\x01J\x01'
    text_for_cjk = text.replace('[J]', marker_J)
    # 去除 CJK 之间的空格（迭代两次确保彻底）
    for _ in range(2):
        text_for_cjk = re.sub(r'([\u4e00-\u9fff])\s+([\u4e00-\u9fff])', r'\1\2', text_for_cjk)
    # 还原 [J]
    text_for_cjk = text_for_cjk.replace(marker_J, '[J]')

    # 3. 去除英文词之间多余空格
    text_for_cjk = re.sub(r'([A-Za-z])\s+([A-Za-z])', r'\1\2', text_for_cjk)

    # 4. 修复 OCR 常见错误
    for old, new in [
        ('\u2014', '-'), ('\u2013', '-'), ('\u203E', '-'), ('~', '-'),
        ('\u3002 ', '\u3002'), ('\uff0c ', '\uff0c'), ('\uff1a ', '\uff1a'), ('\uff1b ', '\uff1b'),
    ]:
        text_for_cjk = text_for_cjk.replace(old, new)

    # 5. 修复断开的多行条目
    text_for_cjk = re.sub(r'(\d)\s*,\s*(\d{4})', r'\1, \2', text_for_cjk)
    text_for_cjk = re.sub(r'(\d)\s*\(\s*(\d+)\s*\)', r'\1(\2)', text_for_cjk)
    text_for_cjk = re.sub(r':\s*(\d)\s*[-]\s*(\d)\b', r': \1-\2', text_for_cjk)

    # 6. 合并被错误断行的条目
    lines = text_for_cjk.split('\n')
    merged = []
    NO_MERGE_END = set('\u3002\uff1b\uff01\uff1f\u2014\u201d\u300d\u300e\u301e\uff1b\u3001\uff0c?)}\"')
    for line in lines:
        line = line.strip()
        if not line:
            merged.append('')
            continue
        if (merged and merged[-1]
                and merged[-1][-1] not in NO_MERGE_END
                and re.match(r'^[A-Za-z0-9(]', line)):
            merged[-1] += ' ' + line
        else:
            merged.append(line)
    text_for_cjk = '\n'.join(merged)

    # 7. 清理多余空格
    text_for_cjk = re.sub(r'\s{2,}', ' ', text_for_cjk)
    return text_for_cjk.strip()



def normalize_ocr_text(text):
    """标准化 OCR 文本，去除 CJK 字符间随机空格，修复常见 OCR 错误"""
    if not text:
        return text

    # 1. 处理 [J] [1] [A] 等标记中的多余空格
    text = re.sub(r'\[\s*J\s*\]', '[J]', text)
    text = re.sub(r'\[\s*(\d+)\s*\]', r'[\1]', text)
    text = re.sub(r'\[\s*([A-Za-z])\s*\]', r'[\1]', text)

    # 2. 去除 CJK 字符之间的空格
    # 用占位符保护 [J]
    marker_J = '\x01J\x01'
    text_for_cjk = text.replace('[J]', marker_J)
    # 去除 CJK 之间的空格（迭代两次确保彻底）
    for _ in range(2):
        text_for_cjk = re.sub(r'([\u4e00-\u9fff])\s+([\u4e00-\u9fff])', r'\1\2', text_for_cjk)
    # 还原 [J]
    text_for_cjk = text_for_cjk.replace(marker_J, '[J]')

    # 3. 去除英文词之间多余空格
    text_for_cjk = re.sub(r'([A-Za-z])\s+([A-Za-z])', r'\1\2', text_for_cjk)

    # 4. 修复 OCR 常见错误
    for old, new in [
        ('\u2014', '-'), ('\u2013', '-'), ('\u203E', '-'), ('~', '-'),
        ('\u3002 ', '\u3002'), ('\uff0c ', '\uff0c'), ('\uff1a ', '\uff1a'), ('\uff1b ', '\uff1b'),
    ]:
        text_for_cjk = text_for_cjk.replace(old, new)

    # 5. 修复断开的多行条目
    text_for_cjk = re.sub(r'(\d)\s*,\s*(\d{4})', r'\1, \2', text_for_cjk)
    text_for_cjk = re.sub(r'(\d)\s*\(\s*(\d+)\s*\)', r'\1(\2)', text_for_cjk)
    text_for_cjk = re.sub(r':\s*(\d)\s*[-]\s*(\d)\b', r': \1-\2', text_for_cjk)

    # 6. 合并被错误断行的条目
    lines = text_for_cjk.split('\n')
    merged = []
    NO_MERGE_END = set('\u3002\uff1b\uff01\uff1f\u2014\u201d\u300d\u300e\u301e\uff1b\u3001\uff0c?)}\"')
    for line in lines:
        line = line.strip()
        if not line:
            merged.append('')
            continue
        if (merged and merged[-1]
                and merged[-1][-1] not in NO_MERGE_END
                and re.match(r'^[A-Za-z0-9(]', line)):
            merged[-1] += ' ' + line
        else:
            merged.append(line)
    text_for_cjk = '\n'.join(merged)

    # 7. 清理多余空格
    text_for_cjk = re.sub(r'\s{2,}', ' ', text_for_cjk)
    return text_for_cjk.strip()



def parse_papers(text):
    """从 OCR 文本解析论文。先按 [J] 标记分段，再按编号分段。"""
    text = normalize_ocr_text(text)

    # 策略A：先找所有 [J] . 期刊名 模式，以它为锚点提取完整块
    # 每个 [J] . xxx , YYYY 结构 = 一篇论文
    j_anchor_pat = re.compile(
        r'('                      # 组1：完整论文块
        r'[\[\(]?\d+[\]\)]?\s*[.、，,、]?\s*'  # 编号
        r'[^\[J]{0,120}'          # 作者+标题（跳过 [J] 区域）
        r'(?:\[[J\]]\s*[.,]\s*)?'  # 可选 [J] 标记
        r'[^,，]{2,60}\s*,\s*20\d\d'  # 期刊名 + 年
        r'[^)]{0,80}'            # 卷期页
        r')'
    )
    j_blocks = [m.group(1).strip() for m in j_anchor_pat.finditer(text) if len(m.group(1)) > 20]

    # 策略B：按编号分段
    raw_lines = [l.strip() for l in text.split('\n') if l.strip()]
    PAPER_START = re.compile(
        r'^(\[\d+\]|\([0-9]+\)|\d+[．、\.、]\s|【\d+】)\s*'
    )
    SECTION_HDR = re.compile(
        r'^[一二三四五六七八九十]+[、．.·]\s*(?:发表|论文|学术|科研|课题|主持|著作|获奖|专利)'
        r'|^(?:期刊论文|学术论文|论著|代表作)[:：]'
    )
    blocks = []
    current = []
    for line in raw_lines:
        if len(line) < 4:
            continue
        if SECTION_HDR.match(line):
            if current:
                blocks.append(' '.join(current))
                current = []
            continue
        if PAPER_START.match(line):
            if current:
                blocks.append(' '.join(current))
            current = [line]
        else:
            current.append(line)
    if current:
        blocks.append(' '.join(current))

    # 合并两种策略，去重
    all_blocks = blocks + j_blocks
    seen = set()
    result = []
    for b in all_blocks:
        b = b.strip()
        if len(b) < 15:
            continue
        key = re.sub(r'\s+', '', b[:60])
        if key in seen:
            continue
        seen.add(key)
        has_year = bool(re.search(r'20\d\d', b))
        has_jmark = bool(re.search(r'[《》[J].]|期刊|学报|杂志', b))
        has_page = bool(re.search(r'\d+\s*[-]\s*\d+', b))
        if has_year and (has_jmark or has_page):
            result.append({'raw': b, 'parts': [b]})

    return result


def extract_paper_info(paper):
    """从论文文本块提取信息 - 期刊名精确版"""
    raw = paper.get('raw', '')
    t = raw

    journal = ''

    # 策略1：精确找 [J] . 期刊名 , 年
    m = re.search(r'\[J\]\s*\.\s*([\u4e00-\u9fffA-Za-z][^\s,，]{0,60})\s*,\s*20\d\d', t)
    if m:
        cand = re.sub(r'\s+', '', m.group(1))
        # 去掉前导噪声（[数字]、作者名等）
        cand = re.sub(r'^\d+\]?\s*', '', cand)  # 去掉 [1] 等
        cand = re.sub(r'^[^\u4e00-\u9fffA-Za-z]+', '', cand)  # 去掉前导非CJK/Latin
        cand = re.sub(r'^[．.、,;：:：\d]+', '', cand)
        # 去掉结尾噪声
        cand = re.sub(r'[．.、,;：:;）]+$', '', cand)
        if 2 < len(cand) < 70:
            journal = cand

    # 策略2：《》
    if not journal:
        m = re.search(r'《([^》]{2,50})》', t)
        if m:
            journal = m.group(1).strip()

    # 策略3：精确找 . 期刊名 , 年（期刊名前是点号）
    if not journal:
        m = re.search(r'\.\s*([\u4e00-\u9fffA-Za-z][^\s,，]{2,60})\s*,\s*20\d\d', t)
        if m:
            cand = re.sub(r'\s+', '', m.group(1))
            cand = re.sub(r'^[^\u4e00-\u9fffA-Za-z\d]+', '', cand)
            cand = re.sub(r'^[．.、,;：:：\d]+', '', cand)
            cand = re.sub(r'[．.、,;：:;）]+$', '', cand)
            if 2 < len(cand) < 70:
                journal = cand

    # ========== 年份 ==========
    year = ''
    for pat in [r'(20[12]\d)年?', r'[\[(](20[12]\d)[)\]]', r',\s*(20[12]\d)[,\s]']:
        m = re.search(pat, t)
        if m:
            yr = re.sub(r'\D', '', m.group(1))
            if 2000 <= int(yr[:4]) <= 2030:
                year = yr[:4]
                break

    # ========== 卷、期 ==========
    vol = ''; issue = ''
    m_vi = re.search(r'(\d{1,4})\s*\(\s*(\d{1,2})\s*\)', t)
    if m_vi:
        vol = m_vi.group(1)
        issue = m_vi.group(2)
    else:
        m_vol = re.search(r'[卷v][olume]*[：.\s]*(\d{1,4})', t, re.I)
        if m_vol:
            vol = m_vol.group(1)
        m_iss = re.search(r'第\s*(\d{1,2})\s*期', t)
        if m_iss:
            issue = m_iss.group(1)

    # ========== 页码 ==========
    pages = ''
    for pat in [r'(\d{1,5})\s*[-]\s*(\d{1,5})']:
        m = re.search(pat, t)
        if m:
            p1, p2 = m.group(1), m.group(2)
            if len(p1) <= 5 and len(p2) <= 5:
                pages = f"{p1}-{p2}"
                break

    # ========== 标题 ==========
    title = ''
    for pat in [r'《([^》]{5,80})》', r'^([A-Z][A-Za-z\s:\-]{10,120})$']:
        m = re.search(pat, t, re.I | re.MULTILINE)
        if m:
            cand = m.group(1).strip()
            if 5 < len(cand) < 120:
                title = cand
                break

    # ========== 作者位置 ==========
    is_first = bool(re.search(r'第一作者|排名第一|共一|共同第一|首位作者', t))
    is_corr = bool(re.search(r'通讯作者|通信作者|Corresponding|联系作者', t, re.I))
    has_marker = bool(re.search(r'作者|author', t, re.I))
    position = []
    if is_first: position.append('第一作者')
    if is_corr: position.append('通讯作者')

    return {
        'journal_name': journal,
        'title': title,
        'year': year,
        'volume': vol,
        'issue': issue,
        'pages': pages,
        'raw': raw[:300],
        'is_valid': bool(position) or has_marker,
        'position': '、'.join(position) if position else '（未标注作者位置）',
    }


# Test
tests = [
    (" . 中华医学杂志 , 2023 , 45(3) : 123 - 130 . （第一作者）", "中华医学杂志/2023"),
    (" . 计算机学报 , 2022 , 38(5) : 456 - 470 . （通讯作者） 3 . 陈六 . 大数据分析方法 [J] . 信息科学学报 , 2024 , 29(2) : 78 - 90 .", "计算机学报/2022"),
    ("[1] 张三 . 人工智能在医学中的应用研究 [J] . 中华医学杂志 , 2023 , 45(3) : 123 - 130 . （第一作者）", "中华医学杂志/2023"),
    ("1. 张三, 李四. 人工智能在医学中的应用研究[J]. 中华医学杂志, 2023, 45(3): 123-130. （第一作者）", "中华医学杂志/2023"),
    ("2. 王五, 赵六. 机器学习算法综述[J]. 计算机学报, 2022, 38(5): 456-470. （通讯作者）", "计算机学报/2022"),
    ("3. 陈八. 深度学习在影像诊断中的应用[J]. 中华放射学杂志, 2024, 58(1): 34-40. 第一作者", "中华放射学杂志/2024"),
]

for t, expected in tests:
    info = extract_paper_info({'raw': t})
    ok = "✓" if info['journal_name'] in expected and info['year'] in expected else "✗"
    print(f"{ok} 期刊:[{info['journal_name']}] 年:{info['year']} 卷:{info['volume']} 期:{info['issue']} 页:{info['pages']} 位置:{info['position']}")

def main():
    st.title("📖 论文情况评价 & 高水平期刊匹配系统")
    st.caption("哈尔滨医科大学人事处 · 仅统计第一作者和通讯作者")

    name_map, issn_map = load_journals()

    with st.sidebar:
        st.header("📋 输入论文信息")
        input_mode = st.radio("输入方式", ["粘贴文本", "上传文件"], horizontal=True)

        text_input = ""
        if input_mode == "粘贴文本":
            text_input = st.text_area(
                "粘贴简历中的论文列表",
                height=400,
                placeholder="支持格式示例：\n1. 张三, 李四. 论文标题[J]. 期刊名, 2024, 35(2): 123-130.（第一作者）\n2. 王五, 等. 论文标题[J]. 期刊名, 2023.（通讯作者）\n3. [1] 赵六. 论文标题[J]. 期刊名, 2022.（第一作者）"
            )
        else:
            uploaded = st.file_uploader(
                "上传简历文件",
                type=['txt', 'docx', 'pdf'],
                help="支持 txt / docx / pdf（含扫描件）"
            )
            if uploaded:
                try:
                    file_bytes = uploaded.read()
                    if uploaded.name.endswith('.docx'):
                        from docx import Document
                        import tempfile, os
                        # docx 需要临时文件
                        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
                        tmp.write(file_bytes)
                        tmp.close()
                        doc = Document(tmp.name)
                        os.unlink(tmp.name)
                        text_input = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                    elif uploaded.name.endswith('.pdf'):
                        # ===== 一站式自动提取：自动尝试所有方法，取最优结果 =====
                        st.info("📄 PDF 处理中，自动选择最佳提取方案...")

                        all_results = {}

                        # 方法1：PyMuPDF blocks（阅读顺序，多栏友好）
                        try:
                            import fitz
                            tmp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                            tmp.write(file_bytes); tmp.close()
                            with fitz.open(tmp.name) as doc:
                                parts = []
                                for page in doc:
                                    blocks = page.get_text("blocks")
                                    blocks.sort(key=lambda b: (round(b[1]/20)*20, b[0]))
                                    page_text = "\n".join([b[4].strip() for b in blocks if b[4].strip()])
                                    if page_text:
                                        parts.append(page_text)
                                txt = "\n".join(parts)
                                all_results["PyMuPDF"] = (txt, len(txt.strip()))
                            os.unlink(tmp.name)
                        except Exception as e:
                            sys.stderr.write("PyMuPDF failed: " + str(e) + "\n")

                        # 方法2：pdfplumber
                        try:
                            import pdfplumber
                            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                                parts = [page.extract_text() or "" for page in pdf.pages if page.extract_text()]
                                txt = "\n".join(parts)
                                all_results["pdfplumber"] = (txt, len(txt.strip()))
                        except Exception as e:
                            sys.stderr.write("pdfplumber failed: " + str(e) + "\n")

                        # 方法3：Tesseract OCR（扫描件）
                        try:
                            import fitz as fitzlib, pytesseract
                            from PIL import Image, ImageEnhance, ImageFilter
                            tmp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                            tmp.write(file_bytes); tmp.close()
                            with fitzlib.open(tmp.name) as doc:
                                ocr_parts = []
                                for page in doc:
                                    mat = fitzlib.Matrix(2, 2)
                                    pix = page.get_pixmap(matrix=mat)
                                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                                    gray = img.convert('L')
                                    img_p = ImageEnhance.Contrast(gray).enhance(1.5).filter(ImageFilter.SHARPEN)
                                    ocr_text = pytesseract.image_to_string(img_p, lang='chi_sim+eng', config='--psm 3 --oem 3')
                                    if ocr_text.strip():
                                        ocr_parts.append(ocr_text.strip())
                                txt = "\n".join(ocr_parts)
                                all_results["TesseractOCR"] = (txt, len(txt.strip()))
                            os.unlink(tmp.name)
                        except Exception as e:
                            sys.stderr.write("Tesseract OCR failed: " + str(e) + "\n")

                        # 自动选择：字符数最多的方法
                        best_name = "PyMuPDF"
                        best_len = 0
                        best_text = ""
                        for name, (txt, clen) in all_results.items():
                            if clen > best_len:
                                best_len = clen
                                best_text = txt
                                best_name = name

                        text_input = best_text
                        strategy_info = {name: clen for name, (txt, clen) in all_results.items()}

                        if best_len == 0:
                            st.error("❌ 无法从 PDF 中提取文字，可能是扫描件或加密 PDF")
                            st.markdown("**建议：** 1. 在 PDF 阅读器中手动全选（Ctrl+A）→ 复制（Ctrl+C）\n2. 切换「粘贴文本」模式，Ctrl+V 粘贴，准确率接近 100%")
                        elif best_len < 200:
                            st.warning("⚠️ 提取文字较少（" + str(best_len) + " 字），可能是扫描件或纯图片 PDF")
                            st.markdown("**建议：** 在 PDF 阅读器中复制文字，粘贴到「粘贴文本」模式")
                        else:
                            st.success("✅ 提取完成（" + best_name + "，" + str(best_len) + " 字）")

                        # 调试面板
                        with st.expander("🔧 调试：各方法提取结果"):
                            for name, (txt, clen) in sorted(all_results.items(), key=lambda x: -x[1][1]):
                                st.caption(name + ": " + str(clen) + " 字")
                            st.text_area("当前文本", value=text_input or "(无)", height=200, key="raw_text_view", label_visibility="collapsed")

                        # 规范化
                        if 'normalize_ocr_text' in globals():
                            text_input = normalize_ocr_text(text_input)
                    else:
                        text_input = file_bytes.decode('utf-8', errors='ignore')
                except Exception as e:
                    st.error(f"读取失败: {e}")

        analyze_btn = st.button("🔍 开始分析", type="primary", use_container_width=True)

        st.divider()
        st.markdown("**高水平期刊目录说明**")
        st.markdown("""
        - 🔴 **高水平1类**：国际顶级科技期刊（CNS级别）
        - 🟠 **高水平2类**：国际一流科技期刊
        - 🟡 **高水平3类**：国际高水平科学期刊
        - 🟢 **高水平4类**：国际重要科技期刊
        - 🔵 **高水平5类**：具有国际影响力国内期刊
        """)

    if not analyze_btn and not text_input:
        st.info("👈 请在左侧输入论文信息后点击「开始分析」")
        return
    if not text_input.strip():
        st.warning("请输入论文内容")
        return

    with st.spinner("正在解析与匹配期刊..."):
        session = requests.Session()
        session.headers.update({'User-Agent': 'Mozilla/5.0 (compatible; PaperEvalBot/1.0)'})

        papers_raw = parse_papers(text_input)
        parsed = [extract_paper_info(p) for p in papers_raw]

        results = []
        session_count = 0
        for p in parsed:
            if not p['is_valid']:
                continue
            jname = p['journal_name']
            matched = match_journal(jname, name_map, issn_map, session) if jname else None
            sci_info = None
            if not matched and jname and session_count < 20:
                session_count += 1
                time.sleep(0.3)
                sci_info = get_sci_zone_from_crossref(jname, session)
            tier = matched['tier'] if matched else None
            tier_level = matched['tier_level'] if matched else 99
            # 组合论文出处信息
            pub_info = ''
            if p.get('year') or p.get('volume') or p.get('issue') or p.get('pages'):
                parts = [p.get('year',''), p.get('volume',''), p.get('issue',''), p.get('pages','')]
                pub_info = ', '.join(x for x in parts if x)

            results.append({
                '论文标题': p.get('title','') or '（未识别标题）',
                '期刊名称': jname if jname else '（未识别）',
                '作者位置': p['position'],
                '目录认定': tier if tier else ('SCI期刊（待核实）' if sci_info else '非目标期刊'),
                '目录级别': tier_level,
                'ISSN': matched.get('issn','') if matched else '',
                'eISSN': matched.get('eissn','') if matched else '',
                '中科院分区': (matched.get('zone','')+'区') if matched and matched.get('zone') else '',
                '学科': matched.get('subject','') if matched else '',
                '发表信息': pub_info,
                '原文': p['raw'][:200],
            })

    if not results:
        st.warning("未识别到有效的第一作者/通讯作者论文，请检查输入格式")
        return

    results.sort(key=lambda x: (x['目录级别'], -len(x['期刊名称'])))
    df = pd.DataFrame(results)

    # ===== 统计概览 =====
    st.divider()
    col1, col2, col3, col4 = st.columns(4)

    total = len(df)
    in_list = len(df[df['目录级别'] < 99])
    first_author = len(df[df['作者位置'].str.contains('第一作者')])
    corresponding = len(df[df['作者位置'].str.contains('通讯作者')])

    col1.metric("有效论文（第一/通讯）", f"{total} 篇")
    col2.metric("高水平目录匹配", f"{in_list} 篇")
    col3.metric("第一作者", f"{first_author} 篇")
    col4.metric("通讯作者", f"{corresponding} 篇")

    # ===== 高水平目录匹配明细（逐条详细列出） =====
    matched_rows = df[df['目录级别'] < 99].to_dict('records')
    if matched_rows:
        st.subheader(f"✅ 高水平期刊目录匹配结果（共 {len(matched_rows)} 篇）")

        tier_order = ['高水平1类','高水平2类','高水平3类','高水平4类',
                      '领军期刊（高水平4类）','高水平5A类','高水平5类',
                      '重点期刊（高水平5类）','高水平5B类']
        tier_colors = {
            '高水平1类':       ('🔴','#ffe5e5'),
            '高水平2类':       ('🟠','#fff3e0'),
            '高水平3类':       ('🟡','#fffde7'),
            '高水平4类':       ('🟢','#e8f5e9'),
            '领军期刊（高水平4类）': ('🟢','#e8f5e9'),
            '高水平5A类':      ('🔵','#e3f2fd'),
            '高水平5类':        ('🔵','#e3f2fd'),
            '重点期刊（高水平5类）': ('🔵','#cce5ff'),
            '高水平5B类':       ('🔵','#f0f8ff'),
        }

        # 按层级分组展示
        for tier in tier_order:
            tier_rows = [r for r in matched_rows if r['目录认定'] == tier]
            if not tier_rows:
                continue
            emoji, bg_color = tier_colors.get(tier, ('⚪','#ffffff'))

            with st.container():
                st.markdown(f"#### {emoji} {tier}（{len(tier_rows)} 篇）")
                for i, row in enumerate(tier_rows, 1):
                    # 提取期刊信息
                    journal_info = row.get('期刊名称', '')
                    pos = row.get('作者位置', '')
                    original_text = row.get('原文', row.get('期刊名称',''))
                    issn_val = row.get('ISSN', '')
                    zone_val = row.get('中科院分区', '')

                    title_val = row.get('论文标题', '')
                    issn_val = row.get('ISSN', '')
                    eissn_val = row.get('eISSN', '')
                    zone_val = row.get('中科院分区', '')
                    subject_val = row.get('学科', '')
                    pub_info = row.get('发表信息', '')
                    original_text = row.get('原文', journal_info)

                    issn_display = issn_val
                    if eissn_val and eissn_val != issn_val:
                        issn_display = f"{issn_val} / {eissn_val}"
                    elif eissn_val:
                        issn_display = eissn_val

                    card_md = f"""
                    <div style="background:{bg_color}; border-radius:8px; padding:14px 18px; 
                                 margin:8px 0; border-left:4px solid #1e3a5f;">
                    <table style="width:100%; border-collapse:collapse; font-size:14px;">
                    <tr>
                        <td style="width:95px; color:#555; font-weight:bold; vertical-align:top;">篇序</td>
                        <td style="font-weight:bold; font-size:15px; color:#1e3a5f;">第{i}篇</td>
                    </tr>
                    <tr>
                        <td style="color:#555; font-weight:bold; vertical-align:top;">论文标题</td>
                        <td><i>{title_val}</i></td>
                    </tr>
                    <tr>
                        <td style="color:#555; font-weight:bold; vertical-align:top;">期刊名称</td>
                        <td><b style="font-size:15px;">{journal_info}</b></td>
                    </tr>
                    <tr>
                        <td style="color:#555; font-weight:bold; vertical-align:top;">目录级别</td>
                        <td><b>{tier}</b></td>
                    </tr>
                    <tr>
                        <td style="color:#555; font-weight:bold; vertical-align:top;">作者位置</td>
                        <td><b>{pos}</b></td>
                    </tr>
                    <tr>
                        <td style="color:#555; font-weight:bold; vertical-align:top;">ISSN</td>
                        <td>{issn_display if issn_display else '—'}</td>
                    </tr>
                    <tr>
                        <td style="color:#555; font-weight:bold; vertical-align:top;">中科院分区</td>
                        <td>{zone_val if zone_val else '—'}{' | ' + subject_val if subject_val else ''}</td>
                    </tr>
                    <tr>
                        <td style="color:#555; font-weight:bold; vertical-align:top;">发表信息</td>
                        <td>{pub_info if pub_info else '—'}</td>
                    </tr>
                    <tr>
                        <td style="color:#555; font-weight:bold; vertical-align:top;">原始条目</td>
                        <td style="color:#555; font-size:12px;">{original_text}</td>
                    </tr>
                    </table>
                    </div>
                    """
                    st.markdown(card_md, unsafe_allow_html=True)
    else:
        st.info("无匹配到高水平目录的论文")

    # ===== 全部论文明细（不含高水平目录的） =====
    other_rows = df[df['目录级别'] >= 99].to_dict('records')
    if other_rows:
        st.subheader(f"⚠️ 非目标期刊 / 待核实（共 {len(other_rows)} 篇）")
        for i, row in enumerate(other_rows, 1):
            journal_info = row.get('期刊名称', '')
            pos = row.get('作者位置', '')
            tier_label = row.get('目录认定', '')
            original_text = row.get('原文', '')
            with st.container():
                title_val = row.get('论文标题', '')
                pub_info = row.get('发表信息', '')
                card_md = f"""
                <div style="background:#f8f8f8; border-radius:8px; padding:12px 18px;
                             margin:6px 0; border-left:3px solid #ccc;">
                <b style="font-size:14px;">❓ 第{i}篇</b> &nbsp;
                <b>{journal_info}</b> &nbsp;|&nbsp; <b>{pos}</b>
                {f'<br><i style="color:#444;font-size:13px;">{title_val}</i>' if title_val else ''}
                {f'<br><span style="color:#555;font-size:12px;">{pub_info}</span>' if pub_info else ''}
                <br><span style="color:#888; font-size:12px;">原始：{original_text[:100]}</span>
                <br><span style="color:#888; font-size:12px;">认定：{tier_label}</span>
                </div>
                """
                st.markdown(card_md, unsafe_allow_html=True)

    # ===== 全部论文明细表 =====
    with st.expander("📋 全部论文明细（表格视图）"):
        disp_df = df[['期刊名称','作者位置','目录认定','目录级别']].copy()
        disp_df.columns = ['期刊名称','作者位置','目录级别','级别值']
        disp_df = disp_df.sort_values('级别值')
        st.dataframe(
            disp_df.style.apply(
                lambda row: [f'background-color: {TIER_BG.get(row["目录级别"], "#fff")}'] * len(row),
                axis=1
            ),
            use_container_width=True,
            hide_index=True
        )

    st.divider()
    st.caption("⚠️ 本系统结果仅供参考，最终认定以学校学术委员会评审为准")

if __name__ == "__main__":
    main()
